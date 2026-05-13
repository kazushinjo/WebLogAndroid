from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Hiragino Kaku Gothic ProN'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')

def set_font(run):
    run.font.name = 'Hiragino Kaku Gothic ProN'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    set_font(run)
    return p

def add_para(text, bold=False, size=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run)
    run.font.size = Pt(11)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, text in enumerate(row):
            cell = r.cells[i]
            cell.text = text
            for run in cell.paragraphs[0].runs:
                set_font(run); run.font.size = Pt(10)
    return t

# ===== タイトル =====
title = doc.add_heading('WebLog for Android — 操作説明書', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_font(run)

add_para('バージョン 1.08 / リリース v0.1.2（2026-05-13 更新）　対応機種: Android 8.0 (API 26) 以上')
doc.add_paragraph()

# ===== 目次 =====
add_heading('目次', 1)
for item in [
    '1. 画面の見かた',
    '2. マイコールサインの登録',
    '3. 交信を記録する',
    '4. 交信記録を検索・絞り込む',
    '5. 交信記録を編集・削除する',
    '6. エクスポート（書き出し）',
    '7. インポート（読み込み）',
    '8. バックアップとリストア',
    '9. 統計情報',
    '10. よくある質問',
    '11. 更新履歴',
]:
    add_bullet(item)

doc.add_page_break()

# ===== 1. 画面の見かた =====
add_heading('1. 画面の見かた', 1)

add_heading('スマートフォン（縦向き）', 2)
add_para('画面上部のタブバーで「登録」と「記録一覧」を切り替えます。')
add_para('【登録タブ】 — QSO（交信記録）の入力フォーム', indent=True)
add_para('【記録一覧タブ】 — 記録済みの交信一覧・検索・統計', indent=True)
add_para('右上の ↑↓ ボタン — インポート・エクスポート・バックアップ画面を開く', indent=True)

doc.add_paragraph()
add_heading('タブレット・横向き（画面幅 600dp 以上）', 2)
add_para('左右2画面で表示されます。iPad版と同様のレイアウトです。')
add_para('【左ペイン】 — QSO（交信記録）の入力フォーム', indent=True)
add_para('【右ペイン】 — 記録済みの交信一覧・検索・統計', indent=True)
add_para('右上の ↑↓ ボタン — インポート・エクスポート・バックアップ画面を開く', indent=True)

# ===== 2. マイコールサインの登録 =====
add_heading('2. マイコールサインの登録', 1)
add_para('アプリを最初に起動したら、まずマイコールサインを登録します。')
doc.add_paragraph()
for s in [
    '1. 「自局」欄の右にある ＋ ボタンをタップ',
    '2. 自分のコールサイン（例: JA1XXX）を入力して「追加」',
]:
    add_para(s, indent=True)
doc.add_paragraph()
add_para('複数のコールサインを登録でき、プルダウンで切り替えられます。コールサインごとに交信記録が独立して管理されます。')
doc.add_paragraph()
add_para('【コールサインを削除するには】削除したいコールサインを選択した状態で — ボタンをタップします。※ コールサインを削除しても、その交信記録データは消えません。')

# ===== 3. 交信を記録する =====
add_heading('3. 交信を記録する', 1)
add_heading('入力の流れ', 2)
for s in [
    '1. コールサインを入力（自動的に大文字になります）',
    '2. 日付を入力（YYYYMMDD 形式　例: 20250509）',
    '3. 時刻を入力（HHMM 形式・JST　例: 1023）',
    '   → 時刻フィールド右の時計アイコンをタップすると現在の日時（JST）が自動入力されます',
    '4. 周波数を MHz 単位で入力（例: 7.100）',
    '   → バンドは自動的に判定されます',
    '5. バンドのドロップダウンで確認・修正できます',
    '6. モードをドロップダウンで選択',
    '7. RST送 / RST受を入力',
    '   → モードに応じてデフォルト値が自動でセットされます（下表参照）',
    '8. 名前・QTH・JCC・QSL・備考を必要に応じて入力',
    '9.「入力」ボタンをタップして記録完了',
]:
    add_para(s, indent=True)

doc.add_paragraph()
add_heading('RST のデフォルト値', 2)
add_table(
    ['モード', 'デフォルト RST'],
    [
        ('SSB / FM / AM / SSTV / D-STAR / C4FM / DMR', '59'),
        ('CW / RTTY / FT8 / FT4 / JS8 / PSK31', '599'),
    ]
)

doc.add_paragraph()
add_heading('バンドの自動判定（周波数入力時）', 2)
add_table(
    ['入力周波数の範囲', '判定バンド'],
    [
        ('1.8〜2.0 MHz', '1.9MHz'), ('3.4〜4.0 MHz', '3.5MHz'),
        ('6.9〜7.5 MHz', '7MHz'),   ('10.0〜10.2 MHz', '10MHz'),
        ('14.0〜14.4 MHz', '14MHz'),('18.0〜18.2 MHz', '18MHz'),
        ('21.0〜21.5 MHz', '21MHz'),('24.8〜25.0 MHz', '24MHz'),
        ('28.0〜29.8 MHz', '28MHz'),('50〜54 MHz', '50MHz'),
        ('140〜150 MHz', '144MHz'), ('420〜440 MHz', '430MHz'),
        ('1200〜1300 MHz', '1200MHz'),('2400〜2450 MHz', '2400MHz'),
    ]
)

doc.add_paragraph()
add_heading('コールサインの重複チェック', 2)
add_para('コールサインを入力すると、過去の交信記録と自動的に照合します。同じ局と過去に交信している場合、⚠ マークと交信済みの件数・直近の記録が表示されます。（エラーではなく参考情報です。そのまま記録できます。）')

doc.add_paragraph()
add_heading('コールサイン履歴の候補', 2)
add_para('2文字以上入力すると、過去に交信したコールサインの候補が表示されます。候補をタップすると コールサイン・名前・QTH・JCC が一括で入力されます。')

doc.add_paragraph()
add_heading('QTH・JCC の入力補助', 2)
add_para('QTH フィールドに市区町村名または読み仮名を入力すると、JCC データベース（1,768 件収録）から候補が表示されます。')
add_para('・候補をタップすると QTH（都道府県＋市区町村名）と JCC コードが自動入力されます', indent=True)
add_para('・JCC コード（4〜6桁）を先に入力すると、対応する都道府県＋市区町村名が QTH に自動入力されます', indent=True)
add_para('・例: コード「1021」→「北海道上磯郡上磯町」、「11001」→「東京都港区」', indent=True)

doc.add_paragraph()
add_heading('QSL の選択肢', 2)
add_table(
    ['値', '意味'],
    [
        ('（空欄）', '未定'),
        ('ビューロー', 'ビューロー経由で送付'),
        ('ダイレクト', 'ダイレクトで送付'),
        ('電子QSL', '電子 QSL'),
        ('なし', 'QSL なし'),
    ]
)

# ===== 4. 検索・絞り込み =====
add_heading('4. 交信記録を検索・絞り込む', 1)
add_para('記録一覧の上部にある検索バーとフィルタを使います。')
doc.add_paragraph()
add_table(
    ['機能', '操作方法'],
    [
        ('テキスト検索', '検索バーに文字を入力（コールサイン・名前・QTH・備考が対象）'),
        ('バンドで絞り込む', '「バンド▼」ボタンでバンドを選択'),
        ('モードで絞り込む', '「モード▼」ボタンでモードを選択'),
        ('件数の確認', 'フィルタボタンの右に「N件」と表示されます'),
    ]
)
doc.add_paragraph()
add_para('フィルタは組み合わせて使えます。')

doc.add_paragraph()
add_heading('一覧の表示形式', 2)
add_para('交信記録は表形式（PC版と同じ列構成）で表示されます。横にスワイプすると全列を確認できます。')
add_para('・列: No / 日付 / 時刻 / コールサイン / バンド / 周波数 / モード / RST送 / RST受 / 名前 / QTH / JCC / QSL / 備考 / 編集 / 削除', indent=True)
add_para('・列ヘッダは縦スクロールしてもスティッキー（固定）表示されます', indent=True)
add_para('・全ての行が同期して横スクロールします', indent=True)

doc.add_paragraph()
add_heading('No（通し番号）について', 2)
add_para('各交信記録に対して、古い順に通し番号 (No) が振られます。一覧では新しい順に表示されるため、No は降順に並びます。')
add_para('・最初の交信が No=1、最新の交信が No=（総交信数）', indent=True)
add_para('・検索・フィルタを変更しても番号は変わりません（PC版のレコード番号と同じ感覚）', indent=True)
add_para('・No 番号はエクスポート時の範囲指定にも使えます（第 6 章を参照）', indent=True)

doc.add_paragraph()
add_heading('表示件数について', 2)
add_para('一覧は最新300件まで表示されます。古い記録を確認したいときは、検索やバンド・モードフィルタで絞り込んでください。')
add_para('絞り込み結果も最新の300件まで表示されます。', indent=True)
add_para('画面下部のステータス行に「総交信数: nnnn」「JCC: nn」などの統計が表示されます。', indent=True)

# ===== 5. 編集・削除 =====
add_heading('5. 交信記録を編集・削除する', 1)
add_heading('編集', 2)
add_para('一覧の各行のどこかをタップすると編集シートが開きます。右端の ✏ アイコンでも開けます。内容を修正して「更新」をタップすると保存されます。')
add_heading('削除', 2)
add_para('一覧の各行にある 🗑 ボタンをタップします。確認ダイアログが表示されるので「削除」をタップすると消去されます。')

# ===== 6. エクスポート =====
add_heading('6. エクスポート（書き出し）', 1)
add_para('右上の ↑↓ ボタン → エクスポートセクションから選択します。')
doc.add_paragraph()
add_heading('ADIF エクスポート', 2)
add_para('形式: ADIF 3.1（.adi）')
add_para('用途: 他のログソフト（Ham Radio Deluxe、Logger32 等）との連携', indent=True)
add_para('ファイル名例: JA1XXX_20250509.adi', indent=True)
add_para('※ 日時は JST から UTC に自動変換されます', indent=True)
doc.add_paragraph()
add_heading('CSV エクスポート', 2)
add_para('形式: CSV（.csv）')
add_para('用途: JA 定番ログソフト「HAMLOG」との連携（HAMLOGデータ互換）', indent=True)
add_para('ファイル名例: JA1XXX_20250509.csv', indent=True)
doc.add_paragraph()
add_para('保存先はファイルピッカーで選択できます（Google Drive・ローカルストレージ等）。')

doc.add_paragraph()
add_heading('範囲指定エクスポート', 2)
add_para('ADIF / CSV / JSON エクスポートボタンを押すと、出力する範囲を選択するダイアログが表示されます。')
doc.add_paragraph()
add_table(
    ['指定方法', '内容'],
    [
        ('全件', '現在のコールサインの全交信記録を出力します'),
        ('番号で指定', '一覧の No 列の番号で範囲を指定します。\n例: No 1 〜 100 → 最も古い 100 件を出力'),
        ('日付で指定', 'YYYYMMDD 形式で範囲を指定します。\n例: 20250101 〜 20250131 → 2025年1月の交信を出力'),
    ]
)
doc.add_paragraph()
add_para('ダイアログ内には、現在の設定で出力される件数がリアルタイムでプレビュー表示されます。「エクスポート」をタップすると、ファイル保存先のピッカーが起動します。')
add_para('・番号で指定する場合、開始 No / 終了 No に空白を入れると 0 件と判定されます', indent=True)
add_para('・日付で指定する場合、YYYYMMDD は 8 桁の数字で入力してください', indent=True)
add_para('・開始 > 終了 になっていても自動的に並び替えて処理します', indent=True)

# ===== 7. インポート =====
add_heading('7. インポート（読み込み）', 1)
add_para('右上の ↑↓ ボタン → インポートセクションから選択します。')
doc.add_paragraph()
add_para('注意: インポートは常に「追記」です。既存の記録は消えません。')
doc.add_paragraph()
add_heading('ADIF インポート', 2)
add_para('.adi または .txt 形式のファイルを読み込みます。他のログソフトや HAMLOG から書き出した ADIF を取り込めます。')
doc.add_paragraph()
add_heading('CSV インポート（自動判別）', 2)
add_para('CSV インポートは2種類の形式を自動判別します。')
add_para('・HAMLOG / Turbo HAMLOG / PC版 weblog の CSV — ヘッダ無し列固定順、Shift-JIS / UTF-8 両対応、日付 YY/MM/DD・時刻 HHMMJ（コロン有無どちらも可）', indent=True)
add_para('・WebLog 日本語ヘッダ CSV — 「コールサイン,日付,時刻,…」のヘッダ付き、UTF-8', indent=True)
add_para('インポート完了後、トーストに判別結果（HAMLOG / 日本語ヘッダ）が表示されます。')
doc.add_paragraph()
add_heading('JSON インポート（バックアップ復元）', 2)
add_para('PC版 WebLog のバックアップ JSON（{version, exported, qsos: [...]} 形式）と Android 版のバックアップ JSON（生の QSO 配列）の両方を読み込めます。PC版の「remarks」フィールドは Android 版の「備考」へ自動マッピングされます。')

# ===== 8. バックアップ =====
add_heading('8. バックアップとリストア', 1)
add_heading('バックアップ（JSON 形式）', 2)
for s in [
    '1. ↑↓ ボタン → 「バックアップ (JSON)」',
    '2. ファイルピッカーで保存先を選択（Google Drive 推奨）',
]:
    add_para(s, indent=True)
add_para('ファイル名例: JA1XXX_backup_20250509.json')
doc.add_paragraph()
add_heading('リストア（バックアップから復元）', 2)
for s in [
    '1. ↑↓ ボタン → 「バックアップから復元 (JSON)」',
    '2. バックアップファイルを選択',
]:
    add_para(s, indent=True)
add_para('注意: リストアは既存データへの追記です。同じファイルを2回読み込むと記録が重複します。全件削除してからリストアすることを推奨します。')
doc.add_paragraph()
add_heading('全件削除', 2)
add_para('↑↓ ボタン → 「全件削除」（確認ダイアログあり）')
add_para('⚠ 全件削除は取り消せません。必ず事前にバックアップを取ってください。')

# ===== 9. 統計 =====
add_heading('9. 統計情報', 1)
add_para('記録一覧の下部に統計が表示されます。')
doc.add_paragraph()
add_table(
    ['項目', '内容'],
    [
        ('総交信数', '現在のコールサインの全記録数（絞り込み後の件数）'),
        ('JCC', '交信した JCC 局数（JCC コード入力済みのもの）'),
        ('バンド別件数', '次の行に表示。例）7MHz:50  144MHz:30'),
    ]
)

# ===== 10. FAQ =====
add_heading('10. よくある質問', 1)
for q, a in [
    ('Q. 時刻は UTC と JST どちらで入力しますか？',
     'A. 時計アイコンボタンでは JST が入力されます。手動入力の場合も JST で入力してください。ADIF エクスポート時に自動的に UTC へ変換されます。'),
    ('Q. 重複チェックが表示されても記録できますか？',
     'A. はい。⚠ は参考情報であり、入力を妨げるものではありません。そのまま「入力」ボタンを押して記録できます。'),
    ('Q. JCC サジェストで候補が出ない場合は？',
     'A. JCC コードを直接入力してください。データベースは 1,768 件収録ですが、新設・廃止等により一致しないことがあります。'),
    ('Q. アプリを削除するとデータも消えますか？',
     'A. はい。アプリを削除するとデータも消えます。定期的にバックアップを取ることを推奨します。'),
    ('Q. 複数の端末でデータを共有できますか？',
     'A. 現バージョンではクラウド自動同期には対応していません。バックアップ JSON ファイルを Google Drive 等で手動共有する方法をご利用ください。'),
    ('Q. インポートしたら記録が重複してしまいました。',
     'A. 全件削除してから再度インポートしてください。削除前に必ずバックアップを取ってください。'),
    ('Q. iPhone/iPad 版とデータを共有できますか？',
     'A. バックアップ JSON 形式・ADIF 形式・CSV 形式はいずれも iPhone/iPad 版と互換性があります。Google Drive や iCloud Drive を経由してファイルを共有することで、相互にデータを移行できます。'),
    ('Q. データ数に上限はありますか？',
     'A. データ数の上限はありません。'),
]:
    add_para(q, bold=True)
    add_para(a, indent=True)
    doc.add_paragraph()

# ===== 11. 更新履歴 =====
add_heading('11. 更新履歴', 1)
add_table(
    ['日付', '更新内容'],
    [
        ('2026-05-13 v0.1.0', '初版リリース（Android版）'),
        ('2026-05-13 v0.1.1',
         '・ボタンのコントラスト改善（白文字に修正）\n'
         '・CSV インポートの自動判別化（HAMLOG / 日本語ヘッダ）\n'
         '・HAMLOG CSV インポート対応（Shift-JIS、列固定順、YY/MM/DD、HHMM[:]J）\n'
         '・PC版 WebLog の JSON バックアップ形式に対応（remarks↔備考 マッピング）\n'
         '・交信一覧を表形式（横スクロール、列固定、スティッキーヘッダ）に変更\n'
         '・一覧表示を最新300件に固定\n'
         '・一覧の行タップで編集モードに移行\n'
         '・JCC サジェスト/コード入力時の QTH に「都道府県＋市区町村名」をセット\n'
         '・インポートをバルク INSERT 化（パフォーマンス改善）'),
        ('2026-05-13 v0.1.2',
         '・交信一覧の先頭に No 列（古い順の通し番号）を追加\n'
         '・ADIF / CSV / JSON エクスポートに範囲指定機能を追加\n'
         '  - 全件 / 番号で指定 / 日付で指定 から選択可能\n'
         '  - 対象件数をリアルタイムでプレビュー表示'),
    ]
)

# ===== 著作権 =====
doc.add_page_break()
add_heading('著作権・注意事項', 1)
doc.add_paragraph()
add_para('© 2026 JH1XHX / JA6FUF　All Rights Reserved.', bold=True)
doc.add_paragraph()
for title_text, body_text in [
    ('【著作権について】',
     '本ドキュメントおよびWebLog for Androidアプリケーションに関する著作権は、JH1XHX/JA6FUFに帰属します。著作権は留保します。'),
    ('【改変について】',
     '本アプリおよび本ドキュメントの改変は自由に行えます。'),
    ('【免責事項】',
     '作者自身はビルドしての動作確認は行なっていません。本アプリおよび本ドキュメントの使用により生じたいかなる損害（データ消失・機器の不具合・その他の損害）についても、著作権者は一切の責任を負いません。ご使用はユーザー自身の責任においておこなってください。'),
    ('【内容変更の予告】',
     'アプリのアップデートに伴い、本ドキュメントの内容が予告なく変更される場合があります。'),
    ('【アマチュア無線法令の遵守】',
     '本アプリはアマチュア無線の交信記録を目的としています。ご使用にあたっては、電波法およびその関連法令を遵守してください。'),
]:
    add_para(title_text, bold=True)
    add_para(body_text, indent=True)
    doc.add_paragraph()

# フッター
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WebLog for Android　操作説明書　v1.08　© 2026 JH1XHX/JA6FUF')
run.font.size = Pt(9)
set_font(run)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'WebLog_Android_操作説明書.docx')
doc.save(out)
print(f'保存: {out}')
