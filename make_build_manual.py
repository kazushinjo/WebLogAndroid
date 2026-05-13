from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# ページ余白
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.0)

# デフォルトフォント
style = doc.styles['Normal']
style.font.name = 'Hiragino Kaku Gothic ProN'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')

def set_font(run):
    run.font.name = 'Hiragino Kaku Gothic ProN'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Hiragino Kaku Gothic ProN')

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    set_font(run)
    return p

def add_para(text, bold=False, size=None, indent=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run)
    run.font.size = Pt(11)
    return p

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'📝 {text}')
    set_font(run)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)
    return p

def add_warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'⚠ {text}')
    set_font(run)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, text in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run); run.bold = True; run.font.size = Pt(10)
    for row in rows:
        r = t.add_row()
        for i, text in enumerate(row):
            cell = r.cells[i]
            cell.text = text
            for run in cell.paragraphs[0].runs:
                set_font(run); run.font.size = Pt(10)
    return t

# ===== タイトル =====
title = doc.add_heading('WebLog for Android — ビルド手順書', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    set_font(run)

add_para('バージョン 1.08　対応機種: Android 8.0 (API 26) 以上')
add_para('作成: 2026年5月　© 2026 JH1XHX/JA6FUF')
doc.add_paragraph()

# ===== 目次 =====
add_heading('目次', 1)
for item in [
    '1. 必要なもの',
    '2. ソースファイルの構成と役割',
    '3. 手順1 — Android Studio をインストールする',
    '4. 手順2 — プロジェクトを開く',
    '5. 手順3 — ランチャーアイコンを設定する（任意）',
    '6. 手順4 — 実機でビルドする',
    '7. 手順5 — エミュレータで確認する（実機なしの場合）',
    '8. トラブルシューティング',
    '9. 動作確認チェックリスト',
    '10. 使用ライブラリ一覧',
    '11. 著作権・注意事項',
]:
    add_bullet(item)

doc.add_page_break()

# ===== 1. 必要なもの =====
add_heading('1. 必要なもの', 1)
add_para('以下の環境が揃っていることを確認してください。')
doc.add_paragraph()
add_table(
    ['項目', '要件', '備考'],
    [
        ['PC / Mac', 'Windows 10/11 / macOS 12以降 / Linux', '8GB以上のRAM推奨'],
        ['Android Studio', 'Ladybug (2024.2.1) 以降推奨', 'JDKは同梱されているため別途不要'],
        ['空きディスク容量', '10GB 以上', 'SDKおよびエミュレータ用'],
        ['Android 端末', 'Android 8.0 (API 26) 以上', 'エミュレータでも可'],
        ['USBケーブル', '実機ビルドの場合のみ', 'データ転送対応ケーブルが必要'],
        ['インターネット接続', '初回ビルド時に必要', 'Gradleライブラリのダウンロードに使用'],
    ]
)

doc.add_paragraph()

# ===== 2. ソースファイルの構成 =====
add_heading('2. ソースファイルの構成と役割', 1)
add_para('プロジェクトは Desktop/WebLogAndroid/ に配置されています。以下にファイル構成を示します。')
doc.add_paragraph()

add_heading('ディレクトリ構造', 2)
structure_lines = [
    'WebLogAndroid/',
    '├── app/',
    '│   ├── build.gradle.kts          ← アプリのビルド設定',
    '│   └── src/main/',
    '│       ├── AndroidManifest.xml   ← アプリの基本設定',
    '│       ├── assets/',
    '│       │   └── jcc.json          ← JCC市区町村データ（1,768件）',
    '│       ├── java/com/weblog/android/',
    '│       │   ├── MainActivity.kt   ← アプリ起動点',
    '│       │   ├── data/',
    '│       │   │   ├── QSO.kt        ← 交信記録データモデル（Room）',
    '│       │   │   ├── QSODao.kt     ← データアクセスオブジェクト',
    '│       │   │   ├── QSODatabase.kt← Roomデータベース定義',
    '│       │   │   └── AppViewModel.kt← アプリ状態管理',
    '│       │   ├── ui/',
    '│       │   │   ├── WebLogApp.kt  ← アプリルート（トースト表示）',
    '│       │   │   ├── MainScreen.kt ← メイン画面（レイアウト振り分け）',
    '│       │   │   ├── EntryFormScreen.kt ← QSO入力フォーム',
    '│       │   │   ├── LogListScreen.kt   ← 交信記録一覧',
    '│       │   │   └── ImportExportScreen.kt ← インポート/エクスポート',
    '│       │   └── utils/',
    '│       │       ├── FreqToBand.kt ← 周波数→バンド変換・定数定義',
    '│       │       ├── ADIFParser.kt ← ADIF/CSV 読み書き',
    '│       │       └── JCCData.kt    ← JCC検索ロジック',
    '│       └── res/values/',
    '│           ├── strings.xml       ← 文字列リソース',
    '│           └── themes.xml        ← アプリテーマ',
    '├── gradle/',
    '│   ├── libs.versions.toml        ← ライブラリバージョン一元管理',
    '│   └── wrapper/gradle-wrapper.properties',
    '├── build.gradle.kts              ← ルートビルド設定',
    '└── settings.gradle.kts           ← プロジェクト設定',
]
for line in structure_lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(line)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_paragraph()

add_heading('主要ファイルの詳細', 2)

add_para('QSO.kt — 交信記録データモデル', bold=True)
add_para('Roomライブラリのエンティティとして定義されています。各フィールドはiOSアプリ（SwiftData）と同一の構造です。UUID形式のIDを自動生成し、myCallフィールドでコールサインごとにデータを分離します。', indent=True)
doc.add_paragraph()

add_para('AppViewModel.kt — アプリ状態管理', bold=True)
add_para('AndroidのViewModelクラスを継承しています。DataStorePreferencesを使用してコールサイン設定を永続化します。交信記録はFlowで購読しており、コールサインが変わると自動的に表示が切り替わります。', indent=True)
doc.add_paragraph()

add_para('MainScreen.kt — メイン画面', bold=True)
add_para('画面幅が600dp以上の場合（タブレット・横向きスマートフォン）は左右2ペインのレイアウト、それ未満の場合はタブ切り替えのレイアウトを自動選択します。iPadアプリの NavigationSplitView と同等の機能です。', indent=True)
doc.add_paragraph()

add_para('ADIFParser.kt — ADIF/CSV 読み書き', bold=True)
add_para('ADIF 3.1形式の読み書きに対応しています。日時はJST⇔UTC変換を行います（iOSアプリと完全互換）。CSVはカンマ区切り・ダブルクォート囲みで出力されます。', indent=True)
doc.add_paragraph()

add_para('JCCData.kt — JCC検索', bold=True)
add_para('assets/jcc.json（1,768件）を起動時に読み込み、市区町村名・読み仮名・JCCコードで前方一致・部分一致検索を行います。iOSアプリと同じjcc.jsonを使用しています。', indent=True)
doc.add_paragraph()

# ===== 3. Android Studio =====
add_heading('3. 手順1 — Android Studio をインストールする', 1)

add_heading('ダウンロードとインストール', 2)
for s in [
    '1. ブラウザで https://developer.android.com/studio を開く',
    '2. 「Download Android Studio」ボタンをクリック',
    '3. 使用しているOSに合ったインストーラをダウンロード',
    '4. インストーラを実行してインストール（デフォルト設定で OK）',
    '5. インストール完了後、Android Studio を起動する',
]:
    add_para(s, indent=True)

doc.add_paragraph()
add_heading('初回起動時の SDK セットアップ', 2)
add_para('Android Studio を初めて起動すると「Android Studio Setup Wizard」が表示されます。')
doc.add_paragraph()
for s in [
    '1. 「Next」をクリックしてウィザードを進める',
    '2. Install Type: 「Standard」を選択して「Next」',
    '3. UIテーマ: 好みのテーマを選択して「Next」',
    '4. SDK Components Setup 画面で内容を確認して「Next」',
    '5. ライセンスに同意して「Finish」',
    '6. SDKのダウンロードが開始されます（数分〜十数分かかります）',
    '7. 完了したら「Finish」をクリック',
]:
    add_para(s, indent=True)

doc.add_paragraph()
add_note('ダウンロードには数GB の空き容量が必要です。完了するまでしばらくお待ちください。')
add_note('企業ネットワーク環境でうまくいかない場合は、Android Studio の Settings > Appearance & Behavior > System Settings > HTTP Proxy からプロキシ設定を行ってください。')

doc.add_paragraph()

# ===== 4. プロジェクトを開く =====
add_heading('4. 手順2 — プロジェクトを開く', 1)

add_heading('プロジェクトのオープン', 2)
for s in [
    '1. Android Studio を起動する',
    '2. ウェルカム画面（Welcome to Android Studio）で「Open」をクリック',
    '   （既にプロジェクトが開いている場合は File > Open...）',
    '3. ファイル選択ダイアログで Desktop/WebLogAndroid/ フォルダを選択',
    '4. 「OK」をクリック',
]:
    add_para(s, indent=True)

doc.add_paragraph()
add_heading('Gradle 同期', 2)
add_para('プロジェクトを開くと、右下に「Gradle sync started」と表示され、Gradle の同期が自動的に開始されます。')
doc.add_paragraph()
for s in [
    '・初回は必要なライブラリをインターネットからダウンロードします（数分かかります）',
    '・下部のステータスバーに進行状況が表示されます',
    '・「Gradle sync finished」と表示されれば同期完了です',
    '・赤いエラーが表示された場合はトラブルシューティング（8章）を参照してください',
]:
    add_para(s, indent=True)

doc.add_paragraph()
add_note('local.properties ファイルは Android Studio が自動的に生成します。手動で作成する必要はありません。このファイルには SDK のパスが記録されており、.gitignore に含まれているためバージョン管理の対象外です。')

doc.add_paragraph()
add_heading('プロジェクト構造の確認', 2)
add_para('左側のプロジェクトナビゲータで「Android」ビューを選択し、以下の構造が表示されることを確認してください。')
doc.add_paragraph()
for s in [
    '・app > manifests > AndroidManifest.xml',
    '・app > java > com.weblog.android > data, ui, utils（各パッケージ）',
    '・app > assets > jcc.json',
    '・app > res > values > strings.xml, themes.xml',
    '・Gradle Scripts > build.gradle.kts (Project: WebLog)',
    '・Gradle Scripts > build.gradle.kts (Module: app)',
]:
    add_para(s, indent=True)

doc.add_paragraph()

# ===== 5. アイコン =====
add_heading('5. 手順3 — ランチャーアイコンを設定する（任意）', 1)
add_para('アプリアイコンをカスタマイズしたい場合の手順です。スキップしてもビルドは可能です（デフォルトのアンドロイドアイコンが使用されます）。')
doc.add_paragraph()
for s in [
    '1. プロジェクトナビゲータで app > src > main > res を右クリック',
    '2. 「New > Image Asset」を選択',
    '3. Icon Type: 「Launcher Icons (Adaptive and Legacy)」のまま',
    '4. Source Asset の「Path」欄に使用したい画像ファイルのパスを指定',
    '5. プレビューを確認して「Next」',
    '6. 「Finish」をクリックしてアイコンを生成',
]:
    add_para(s, indent=True)

doc.add_paragraph()
add_note('アイコン画像は 1024×1024px 以上の PNG ファイルを推奨します。背景が透過（アルファチャンネルあり）のものがアダプティブアイコンに適しています。')

doc.add_paragraph()

# ===== 6. 実機ビルド =====
add_heading('6. 手順4 — 実機でビルドする', 1)

add_heading('Android 端末の開発者モードを有効にする', 2)
add_para('まず端末側で「開発者向けオプション」を有効にする必要があります。')
doc.add_paragraph()
for s in [
    '1. 端末の「設定」アプリを開く',
    '2.「デバイス情報」（機種によって「端末情報」「携帯情報」等）をタップ',
    '3.「ビルド番号」を7回連続でタップ',
    '   → 「開発者になりました！」というメッセージが表示される',
    '4. 設定に戻り「開発者向けオプション」をタップ（「システム」の中にある場合あり）',
    '5.「USB デバッグ」をオンにする',
]:
    add_para(s, indent=True)
doc.add_paragraph()
add_warn('端末のメーカーや Android バージョンによって「ビルド番号」の場所が異なります。見つからない場合は「（メーカー名） ビルド番号」で検索してください。')

doc.add_paragraph()
add_heading('PC/Mac への接続', 2)
for s in [
    '1. USB ケーブル（データ転送対応）で端末と PC/Mac を接続',
    '2. 端末に「USB デバッグを許可しますか？」ダイアログが表示される',
    '3.「このコンピュータから常に許可する」にチェックして「OK」をタップ',
    '4. Android Studio 上部のデバイス選択に端末名が表示されることを確認',
]:
    add_para(s, indent=True)
doc.add_paragraph()
add_note('「充電のみ」モードになっている場合はファイル転送（MTP）に切り替えてください。通知バーから「USB 接続の使用目的」を変更できます。')

doc.add_paragraph()
add_heading('ビルドして実行', 2)
for s in [
    '1. Android Studio 上部のデバイス選択ドロップダウンで接続した端末を選択',
    '2.「▶ Run」ボタン（緑の三角）をクリック',
    '   または Shift + F10（Windows/Linux）/ Ctrl + R（Mac）',
    '3. ビルドが開始されます（初回は数十秒〜数分かかります）',
    '4. ビルドが完了すると端末にアプリがインストールされ、自動的に起動します',
    '5. 「WebLog v1.08」と表示されれば成功です',
]:
    add_para(s, indent=True)
doc.add_paragraph()
add_note('ビルドの進行状況は画面下部の「Build」タブで確認できます。「BUILD SUCCESSFUL」が表示されれば正常完了です。')
add_warn('初回インストール時に端末側で「このアプリのインストールを許可しますか？」と表示された場合は「インストール」をタップしてください。')

doc.add_paragraph()

# ===== 7. エミュレータ =====
add_heading('7. 手順5 — エミュレータで確認する（実機なしの場合）', 1)
add_para('実機がない場合、Android Studio に内蔵されたエミュレータ（AVD: Android Virtual Device）でも動作確認が可能です。ただし、ファイルのエクスポート/インポートの動作は実機と異なる場合があります。最終確認は実機を推奨します。')
doc.add_paragraph()

add_heading('エミュレータの作成', 2)
for s in [
    '1. Android Studio 右側の「Device Manager」アイコンをクリック',
    '   （メニューから Tools > Device Manager でも開けます）',
    '2.「Create Device」（または「+」ボタン）をクリック',
    '3. ハードウェアプロファイルの選択:',
    '   ・スマートフォンの場合: Phone カテゴリの「Pixel 8」を選択して「Next」',
    '   ・タブレットの場合: Tablet カテゴリの「Pixel Tablet」を選択して「Next」',
    '4. システムイメージの選択:',
    '   ・「Android 14.0 (API 35)」を選択',
    '   ・ダウンロードが必要な場合は「Download」リンクをクリック',
    '   ・ダウンロード完了後「Next」',
    '5.「Finish」でエミュレータの作成完了',
]:
    add_para(s, indent=True)
doc.add_paragraph()
add_warn('エミュレータのシステムイメージは数GB の空き容量が必要です。')

doc.add_paragraph()
add_heading('エミュレータの起動と実行', 2)
for s in [
    '1. Device Manager の一覧に作成したエミュレータが表示される',
    '2. 「▶ 起動」ボタンをクリックしてエミュレータを起動',
    '3. エミュレータが起動したら、Android Studio 上部のデバイス選択でエミュレータを選択',
    '4.「▶ Run」ボタンでアプリをビルド・インストール',
]:
    add_para(s, indent=True)
doc.add_paragraph()
add_note('タブレット用エミュレータ（Pixel Tablet 等）を使うと、左右2ペインのレイアウトを確認できます。')

doc.add_paragraph()

# ===== 8. トラブルシューティング =====
add_heading('8. トラブルシューティング', 1)

troubles = [
    (
        'Gradle の同期に失敗する',
        ['ネットワーク接続の問題', 'Gradleのキャッシュの不整合', 'プロキシ設定の問題'],
        [
            'Android Studio メニュー → File > Sync Project with Gradle Files を実行',
            '改善しない場合: File > Invalidate Caches > Invalidate and Restart を実行',
            '企業ネットワークの場合: Settings > Appearance & Behavior > System Settings > HTTP Proxy からプロキシを設定',
        ]
    ),
    (
        '「Could not resolve com.android...」エラー',
        ['ライブラリのダウンロードに失敗している', 'ネットワーク接続が不安定'],
        [
            'ネットワーク接続を確認してから再同期',
            'gradle/libs.versions.toml のバージョン番号が最新の Maven リポジトリに存在するか確認',
            'それでも失敗する場合はバージョン番号を少し古いものに変更して試す',
        ]
    ),
    (
        'jcc.json が読み込まれない（JCCサジェストが動作しない）',
        ['assets フォルダにファイルが存在しない', 'ファイルが正しくコピーされていない'],
        [
            'app/src/main/assets/jcc.json が存在するか確認',
            'ない場合は WebLogApp/WebLog/Data/jcc.json を app/src/main/assets/ にコピー',
            'Android Studio で「Sync Project with Gradle Files」を実行後、再ビルド',
        ]
    ),
    (
        '端末が Android Studio に認識されない',
        ['USB デバッグが有効になっていない', 'USB ドライバが未インストール（Windows）', 'データ転送モードになっていない'],
        [
            '端末側で「開発者向けオプション > USB デバッグ」がオンになっているか確認',
            '端末の通知バーから USB 接続の種類を「ファイル転送」に変更',
            'Windows の場合: USB ドライバをメーカーサイトからインストール',
            '別の USB ケーブルや USB ポートを試す',
        ]
    ),
    (
        '「INSTALL_FAILED_VERSION_DOWNGRADE」エラー',
        ['端末に新しいバージョンのアプリが既にインストールされている'],
        [
            '端末の設定でアプリを一度アンインストール（アプリのデータは消えます）',
            '再度「▶ Run」でビルド・インストール',
        ]
    ),
    (
        '「Kotlin version is not compatible」エラー',
        ['KSP プラグインのバージョンと Kotlin バージョンが一致していない'],
        [
            'gradle/libs.versions.toml を開き kotlin と ksp のバージョンを確認',
            'ksp のバージョンは「2.0.21-1.0.27」のように「Kotlinバージョン-KSPバージョン」の形式',
            '最新の対応バージョンは github.com/google/ksp/releases で確認',
        ]
    ),
    (
        'ビルドは成功するがアプリがクラッシュする',
        ['実行時エラーの可能性'],
        [
            'Android Studio 下部の「Logcat」タブを開く',
            '赤字の「E/」で始まるエラーログを確認',
            'よくある原因: jcc.json の読み込み失敗（assets に配置されているか確認）',
        ]
    ),
]

for err, causes, fixes in troubles:
    add_para(f'❌ {err}', bold=True)
    add_para('【原因】', bold=False)
    for c in causes:
        add_para(f'・{c}', indent=True)
    add_para('【対処】', bold=False)
    for f in fixes:
        add_para(f'・{f}', indent=True)
    doc.add_paragraph()

# ===== 9. チェックリスト =====
add_heading('9. 動作確認チェックリスト', 1)
add_para('ビルド後、以下の項目を順番に確認してください。')
doc.add_paragraph()

add_heading('基本動作', 2)
for item in [
    'アプリが起動する（「WebLog v1.08」がタイトルバーに表示される）',
    'マイコールサインを追加できる（＋ボタン → コールサイン入力 → 追加）',
    '追加したコールサインがプルダウンに表示される',
    'コールサインを削除できる（－ボタン）',
]:
    add_bullet(f'□  {item}')

doc.add_paragraph()
add_heading('交信記録の入力', 2)
for item in [
    'QSO を入力して登録できる（コールサイン・日付・時刻・周波数・モードを入力して「入力」ボタン）',
    '「現在」ボタンを押すと現在の日時（JST）が自動入力される',
    '周波数を入力するとバンドが自動判定される（例: 7.100 → 7MHz）',
    'モードを変更するとRSTのデフォルト値が切り替わる（SSB→59、CW→599）',
    'コールサイン入力で過去の交信がある場合に重複チェックが表示される',
    '2文字以上入力するとコールサイン候補が表示される',
    'QTH欄に市区町村名を入力するとJCC候補が表示される',
    'JCC候補をタップするとQTHとJCCコードが自動入力される',
    'JCCコードを直接入力すると市区町村名がQTHに自動入力される',
]:
    add_bullet(f'□  {item}')

doc.add_paragraph()
add_heading('交信記録一覧', 2)
for item in [
    '交信一覧に記録が表示される',
    '検索バーに文字を入力すると絞り込まれる',
    'バンドフィルタで絞り込みができる',
    'モードフィルタで絞り込みができる',
    '✏ ボタンで編集シートが開き、更新できる',
    '🗑 ボタンで確認ダイアログが表示され、削除できる',
    '画面下部に統計情報（総交信数・JCC数・バンド別件数）が表示される',
]:
    add_bullet(f'□  {item}')

doc.add_paragraph()
add_heading('インポート・エクスポート', 2)
for item in [
    'ADIF エクスポートでファイルが書き出せる（.adi）',
    'CSV エクスポートでファイルが書き出せる（.csv）',
    'バックアップ JSON が保存できる（.json）',
    'ADIF インポートで記録を追記できる',
    'CSV インポートで記録を追記できる',
    'バックアップ JSON からリストアできる',
    '全件削除が確認ダイアログ付きで動作する',
]:
    add_bullet(f'□  {item}')

doc.add_paragraph()
add_heading('レイアウト', 2)
for item in [
    'スマートフォン（縦向き）でタブ切り替えが動作する（登録 / 記録一覧）',
    'タブレットまたはスマートフォン横向き（600dp以上）で左右2ペインが表示される',
]:
    add_bullet(f'□  {item}')

doc.add_paragraph()

# ===== 10. 使用ライブラリ =====
add_heading('10. 使用ライブラリ一覧', 1)
add_table(
    ['ライブラリ', 'バージョン', '用途'],
    [
        ['Jetpack Compose BOM', '2024.09.02', 'UIフレームワーク（SwiftUIに相当）'],
        ['Material 3', 'BOM管理', 'マテリアルデザインコンポーネント'],
        ['Material Icons Extended', 'BOM管理', 'アイコンセット'],
        ['Room', '2.6.1', 'データベース（SwiftDataに相当）'],
        ['KSP', '2.0.21-1.0.27', 'Room用コードジェネレータ'],
        ['DataStore Preferences', '1.1.1', '設定値の永続化（UserDefaultsに相当）'],
        ['Lifecycle ViewModel Compose', '2.8.4', 'ViewModelとComposeの連携'],
        ['Activity Compose', '1.9.1', 'ComposeのActivity統合'],
        ['Coroutines Android', '1.8.1', '非同期処理（async/awaitに相当）'],
        ['Gson', '2.11.0', 'JSON読み書き（Codableに相当）'],
    ]
)

doc.add_paragraph()

# ===== 11. 著作権・注意事項 =====
doc.add_page_break()
add_heading('11. 著作権・注意事項', 1)
doc.add_paragraph()
add_para('© 2026 JH1XHX / JA6FUF　All Rights Reserved.', bold=True)
doc.add_paragraph()

for title_text, body_text in [
    ('【著作権について】',
     '本ドキュメントおよびWebLog for Androidアプリケーションに関する著作権は、JH1XHX/JA6FUFに帰属します。'),
    ('【改変について】',
     '本アプリおよび本ドキュメントの改変は自由に行えます。ただし、改変を行った場合は必ず著作権者（JH1XHX/JA6FUF）にご連絡ください。'),
    ('【免責事項】',
     '本アプリおよび本ドキュメントの使用により生じたいかなる損害（データ消失・機器の不具合・その他の損害）についても、著作権者は一切の責任を負いません。ご使用はユーザー自身の責任においておこなってください。'),
    ('【内容変更の予告】',
     'アプリのアップデートに伴い、本ドキュメントの内容が予告なく変更される場合があります。最新情報は配布元をご確認ください。'),
    ('【アマチュア無線法令の遵守】',
     '本アプリはアマチュア無線の交信記録を目的としています。ご使用にあたっては、電波法およびその関連法令を遵守してください。'),
]:
    add_para(title_text, bold=True)
    add_para(body_text, indent=True)
    doc.add_paragraph()

# フッター
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('WebLog for Android　ビルド手順書　v1.08　© 2026 JH1XHX/JA6FUF')
run.font.size = Pt(9)
set_font(run)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'WebLog_Android_ビルド手順書.docx')
doc.save(out)
print(f'保存: {out}')
